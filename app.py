import streamlit as st
import os
import cv2
import json
import datetime
import requests
import tempfile
import numpy as np
from PIL import Image
from io import BytesIO

from azure.ai.vision.imageanalysis import ImageAnalysisClient
from azure.ai.vision.imageanalysis.models import VisualFeatures
from azure.core.credentials import AzureKeyCredential
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# CONFIG – reads from .env locally, Streamlit Secrets in production
# ==========================================
def get_config():
    return {
        "AZURE_KEY":      st.secrets.get("AZURE_KEY",      os.getenv("AZURE_KEY", "")),
        "AZURE_ENDPOINT": st.secrets.get("AZURE_ENDPOINT", os.getenv("AZURE_ENDPOINT", "")),
        "GROQ_API_KEY":   st.secrets.get("GROQ_API_KEY",   os.getenv("GROQ_API_KEY", "")),
        "GROQ_MODEL":     st.secrets.get("GROQ_MODEL",     os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")),
    }

# ==========================================
# MEMORY MANAGER
# ==========================================
MEMORY_FILE = "agent_memory.json"

def load_memory():
    if os.path.exists(MEMORY_FILE):
        with open(MEMORY_FILE, "r") as f:
            return json.load(f)
    return {}

def save_memory(vocab):
    with open(MEMORY_FILE, "w") as f:
        json.dump(vocab, f, indent=2)

def learn_correction(vocab, bad_word, good_word):
    if bad_word.lower() != good_word.lower():
        vocab[bad_word] = good_word
        save_memory(vocab)

def apply_memory_rules(line, vocab):
    words = line.split()
    corrected = []
    for word in words:
        core = word.strip(".,;:()[]{}")
        if core in vocab:
            word = word.replace(core, vocab[core])
        corrected.append(word)
    return " ".join(corrected)

# ==========================================
# AUDIT LOGGER
# ==========================================
LOG_FILE = "agent_audit.log"

def audit_log(event_type, detail):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{timestamp}] [{event_type.upper()}] {detail}\n"
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(line)

def read_log():
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, "r", encoding="utf-8") as f:
            return f.read()
    return "(No log entries yet)"

# ==========================================
# LLM CLEANUP via Groq
# ==========================================
BATCH_SIZE = 20

def call_groq_batch(lines, cfg):
    system_prompt = (
        "You are an expert OCR post-processor for handwritten notes/text. "
        "Fix OCR errors (letter confusion, missing spaces, garbled symbols) "
        "while preserving the original meaning exactly. "
        "Also identify which lines are headings/titles vs. body text. "
        "Respond ONLY with valid JSON — no markdown fences, no extra text:\n"
        '{"cleaned_lines": ["..."], "heading_flags": [true/false, ...]}'
    )
    user_prompt = (
        "Fix OCR errors in these text lines. "
        "Return EVERY line (even unchanged ones) in order:\n\n"
        + "\n".join(f"{i+1}. {l}" for i, l in enumerate(lines))
    )
    response = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {cfg['GROQ_API_KEY']}",
            "Content-Type": "application/json",
        },
        json={
            "model": cfg["GROQ_MODEL"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            "temperature": 0.1,
            "max_tokens": 2048,
        },
        timeout=30,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"].strip()
    content = content.replace("```json", "").replace("```", "").strip()
    return json.loads(content)

def llm_cleanup(raw_text, cfg):
    all_lines   = raw_text.split("\n")
    out_cleaned = []
    out_flags   = []
    errors      = []

    for batch_start in range(0, len(all_lines), BATCH_SIZE):
        batch = all_lines[batch_start : batch_start + BATCH_SIZE]
        try:
            result  = call_groq_batch(batch, cfg)
            cleaned = result.get("cleaned_lines", batch)
            flags   = result.get("heading_flags", [False] * len(batch))
            while len(cleaned) < len(batch):
                cleaned.append(batch[len(cleaned)])
            while len(flags) < len(batch):
                flags.append(False)
            out_cleaned.extend(cleaned[: len(batch)])
            out_flags.extend(flags[: len(batch)])
        except Exception as e:
            errors.append(str(e))
            out_cleaned.extend(batch)
            out_flags.extend([False] * len(batch))

    result = {"cleaned_lines": out_cleaned, "heading_flags": out_flags}
    if errors:
        result["error"] = " | ".join(errors)
    return result

# ==========================================
# AZURE OCR
# ==========================================
def run_azure_ocr(image_bytes, cfg):
    client = ImageAnalysisClient(
        endpoint=cfg["AZURE_ENDPOINT"],
        credential=AzureKeyCredential(cfg["AZURE_KEY"]),
    )
    result = client.analyze(
        image_data=image_bytes,
        visual_features=[VisualFeatures.READ],
    )
    lines = []
    img_width = 1
    if result.read and result.read.blocks:
        for line in result.read.blocks[0].lines:
            xs = [p.x for p in line.bounding_polygon]
            ys = [p.y for p in line.bounding_polygon]
            lines.append({
                "original_text": line.text,
                "agent_text":    line.text,
                "final_text":    line.text,
                "is_heading":    False,
                "flagged":       False,
                "x": min(xs),
                "y": min(ys),
            })
        # estimate width from bounding boxes
        all_xs = [p.x for block in result.read.blocks for line in block.lines for p in line.bounding_polygon]
        img_width = max(all_xs) if all_xs else 1

    lines.sort(key=lambda r: r["y"])
    return lines, img_width

# ==========================================
# WORD DOC GENERATION
# ==========================================
def generate_docx(lines_data, img_width):
    doc = Document()
    for line in lines_data:
        text = line["final_text"].strip()
        if not text:
            continue
        para = doc.add_paragraph()
        if line["is_heading"]:
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
        else:
            para.add_run(text)
        x = line["x"]
        if x > img_width * 0.6:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif img_width * 0.25 < x < img_width * 0.45:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==========================================
# PAGE CONFIG & CUSTOM CSS
# ==========================================
st.set_page_config(
    page_title="Agentic OCR System",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Mono:wght@400;700&family=Syne:wght@400;600;800&display=swap');

html, body, [class*="css"] {
    font-family: 'Syne', sans-serif;
}

.main { background: #0d0d0d; }

.stApp {
    background: linear-gradient(135deg, #0d0d0d 0%, #111827 50%, #0d0d0d 100%);
    color: #e5e5e5;
}

h1, h2, h3 {
    font-family: 'Syne', sans-serif;
    font-weight: 800;
}

.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: 2.8rem;
    font-weight: 800;
    background: linear-gradient(90deg, #00ff88, #00ccff, #a855f7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: -1px;
    line-height: 1.1;
    margin-bottom: 0.2rem;
}

.hero-sub {
    font-family: 'Space Mono', monospace;
    font-size: 0.8rem;
    color: #666;
    letter-spacing: 2px;
    text-transform: uppercase;
    margin-bottom: 2rem;
}

.step-badge {
    background: linear-gradient(135deg, #00ff8820, #00ccff20);
    border: 1px solid #00ff8840;
    border-radius: 8px;
    padding: 0.6rem 1rem;
    font-family: 'Space Mono', monospace;
    font-size: 0.75rem;
    color: #00ff88;
    display: inline-block;
    margin-bottom: 1rem;
    letter-spacing: 1px;
}

.flagged-row {
    background: #2a1f00;
    border-left: 3px solid #f59e0b;
    padding: 4px 8px;
    border-radius: 0 4px 4px 0;
    margin: 2px 0;
    font-family: 'Space Mono', monospace;
    font-size: 0.8rem;
}

.normal-row {
    background: #1a1a2e;
    border-left: 3px solid #00ff8840;
    padding: 4px 8px;
    border-radius: 0 4px 4px 0;
    margin: 2px 0;
    font-family: 'Space Mono', monospace;
    font-size: 0.8rem;
}

.metric-card {
    background: #161b2e;
    border: 1px solid #1e293b;
    border-radius: 12px;
    padding: 1.2rem;
    text-align: center;
}

.metric-num {
    font-family: 'Space Mono', monospace;
    font-size: 2rem;
    font-weight: 700;
    color: #00ff88;
}

.metric-label {
    font-size: 0.75rem;
    color: #666;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-top: 0.2rem;
}

.consent-box {
    background: #1a0000;
    border: 1px solid #ff444440;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1rem 0;
}

.stButton > button {
    font-family: 'Syne', sans-serif;
    font-weight: 600;
    border-radius: 8px;
    border: none;
    transition: all 0.2s;
}

.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 4px 20px rgba(0,255,136,0.2);
}

div[data-testid="stSidebar"] {
    background: #111111;
    border-right: 1px solid #1e293b;
}

.sidebar-section {
    font-family: 'Space Mono', monospace;
    font-size: 0.7rem;
    color: #555;
    text-transform: uppercase;
    letter-spacing: 2px;
    margin: 1rem 0 0.5rem 0;
}
</style>
""", unsafe_allow_html=True)

# ==========================================
# SESSION STATE INIT
# ==========================================
if "ocr_lines"   not in st.session_state: st.session_state.ocr_lines   = []
if "img_width"   not in st.session_state: st.session_state.img_width   = 1
if "step"        not in st.session_state: st.session_state.step        = "upload"
if "vocab"       not in st.session_state: st.session_state.vocab       = load_memory()
if "llm_used"    not in st.session_state: st.session_state.llm_used    = False
if "consented"   not in st.session_state: st.session_state.consented   = False

cfg = get_config()

# ==========================================
# SIDEBAR
# ==========================================
with st.sidebar:
    st.markdown('<div class="hero-title" style="font-size:1.4rem;">⬡ OCR Agent</div>', unsafe_allow_html=True)
    st.markdown('<div class="hero-sub" style="font-size:0.65rem;">Phase 2 · Agentic System</div>', unsafe_allow_html=True)

    st.markdown('<div class="sidebar-section">Configuration</div>', unsafe_allow_html=True)
    use_llm = st.toggle("Groq LLM Cleanup", value=True, help="Send OCR text to Groq for intelligent error correction")

    st.markdown('<div class="sidebar-section">Memory</div>', unsafe_allow_html=True)
    st.caption(f"**{len(st.session_state.vocab)}** corrections learned")

    if st.button("🧠 Manage Memory", use_container_width=True):
        st.session_state.show_memory = not st.session_state.get("show_memory", False)

    if st.session_state.get("show_memory", False):
        if st.session_state.vocab:
            for k, v in list(st.session_state.vocab.items()):
                col1, col2, col3 = st.columns([3, 3, 1])
                col1.caption(f"`{k}`")
                col2.caption(f"→ `{v}`")
                if col3.button("✕", key=f"del_{k}"):
                    del st.session_state.vocab[k]
                    save_memory(st.session_state.vocab)
                    audit_log("MEMORY_DELETE", f"Removed correction for '{k}'")
                    st.rerun()
        else:
            st.caption("No corrections yet.")

    st.markdown('<div class="sidebar-section">Audit Log</div>', unsafe_allow_html=True)
    if st.button("📋 View Log", use_container_width=True):
        st.session_state.show_log = not st.session_state.get("show_log", False)

    if st.session_state.get("show_log", False):
        log_content = read_log()
        st.text_area("", value=log_content, height=200, disabled=True, label_visibility="collapsed")

    st.markdown("---")
    st.markdown('<div class="sidebar-section">Workflow</div>', unsafe_allow_html=True)
    steps = ["upload", "consent", "observe", "review", "done"]
    step_labels = {"upload": "① Upload", "consent": "② Consent", "observe": "③ OCR + LLM", "review": "④ Review", "done": "✓ Complete"}
    for s in steps:
        active = st.session_state.step == s
        color = "#00ff88" if active else "#333"
        st.markdown(f'<div style="color:{color}; font-family:Space Mono; font-size:0.72rem; padding:2px 0;">{step_labels[s]}</div>', unsafe_allow_html=True)

# ==========================================
# MAIN CONTENT
# ==========================================
st.markdown('<div class="hero-title">Agentic OCR System</div>', unsafe_allow_html=True)
st.markdown('<div class="hero-sub">Observe → Interpret → Decide → Act → Learn</div>', unsafe_allow_html=True)

# ── STEP: UPLOAD ──────────────────────────────────────────────────────────────
if st.session_state.step == "upload":
    st.markdown('<div class="step-badge">STEP 01 / UPLOAD IMAGE</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1.2, 1])
    with col1:
        uploaded = st.file_uploader(
            "Drop your handwritten note image here",
            type=["png", "jpg", "jpeg"],
            label_visibility="visible",
        )
        if uploaded:
            st.image(uploaded, caption="Uploaded image", use_container_width=True)
            st.session_state.uploaded_file  = uploaded
            st.session_state.uploaded_name  = uploaded.name

    with col2:
        st.markdown("#### How it works")
        st.markdown("""
        **① Upload** a handwritten note image  
        **② Consent** to external API usage  
        **③ OCR** extracts text via Azure Vision  
        **④ LLM** cleans up errors via Groq  
        **⑤ Review** and correct in a live editor  
        **⑥ Download** your formatted Word doc  
        """)
        st.info("Your corrections are saved locally and improve future scans automatically.", icon="🧠")

    if uploaded:
        if st.button("▶  Continue to Consent →", type="primary", use_container_width=False):
            st.session_state.step = "consent"
            st.rerun()

# ── STEP: CONSENT ─────────────────────────────────────────────────────────────
elif st.session_state.step == "consent":
    st.markdown('<div class="step-badge">STEP 02 / PRIVACY & CONSENT</div>', unsafe_allow_html=True)

    st.markdown('<div class="consent-box">', unsafe_allow_html=True)
    st.markdown("### ⚠ Data Privacy Notice")
    st.markdown("""
This application will send your image to **external APIs**:

- 🔵 **Microsoft Azure Computer Vision** — performs OCR on your image
- 🟣 **Groq Cloud LLM API** — cleans up OCR errors using an LLM

**What is stored:**
- No images or text are retained by Azure or Groq beyond the API call
- Your corrections are saved **locally** in `agent_memory.json`
- All agent actions are logged **locally** in `agent_audit.log`

By proceeding, you consent to this data processing.
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("✅ I Agree — Proceed", type="primary"):
            st.session_state.consented = True
            audit_log("CONSENT", "User granted data processing consent.")
            st.session_state.step = "observe"
            st.rerun()
    with col2:
        if st.button("✕ Cancel"):
            audit_log("CONSENT", "User declined data processing consent.")
            st.session_state.step = "upload"
            st.rerun()

# ── STEP: OCR + LLM ───────────────────────────────────────────────────────────
elif st.session_state.step == "observe":
    st.markdown('<div class="step-badge">STEP 03 / OBSERVE + INTERPRET</div>', unsafe_allow_html=True)

    if not cfg["AZURE_KEY"] or not cfg["AZURE_ENDPOINT"]:
        st.error("Azure credentials are missing. Add them to `.streamlit/secrets.toml`.")
        st.stop()

    with st.spinner("🔍 Running Azure OCR..."):
        try:
            image_bytes = st.session_state.uploaded_file.read()
            lines, img_width = run_azure_ocr(image_bytes, cfg)
            st.session_state.ocr_lines = lines
            st.session_state.img_width = img_width
            audit_log("OBSERVE", f"Azure extracted {len(lines)} lines from {st.session_state.uploaded_name}")

            # Apply memory rules
            for item in lines:
                item["agent_text"] = apply_memory_rules(item["original_text"], st.session_state.vocab)

        except Exception as e:
            audit_log("ERROR", str(e))
            st.error(f"Azure OCR failed: {e}")
            st.stop()

    if use_llm and cfg["GROQ_API_KEY"]:
        with st.spinner("🤖 Groq LLM cleanup in progress..."):
            try:
                raw_block  = "\n".join(r["agent_text"] for r in lines)
                llm_result = llm_cleanup(raw_block, cfg)

                if "error" in llm_result:
                    st.warning(f"LLM partial error (fallback used): {llm_result['error']}")
                    audit_log("LLM_ERROR", llm_result["error"])

                cleaned = llm_result.get("cleaned_lines", [])
                flags   = llm_result.get("heading_flags", [])

                for i, item in enumerate(lines):
                    if i < len(cleaned):
                        llm_line = cleaned[i]
                        if llm_line.strip() != item["agent_text"].strip():
                            item["flagged"] = True
                            audit_log("LLM_CHANGE", f"Line {i+1}: '{item['agent_text']}' → '{llm_line}'")
                        item["agent_text"] = llm_line
                        item["is_heading"] = flags[i] if i < len(flags) else False

                st.session_state.llm_used = True
                audit_log("LLM", "Groq cleanup applied successfully.")

            except Exception as e:
                audit_log("LLM_ERROR", str(e))
                st.warning(f"LLM failed, using rule-only mode: {e}")
    else:
        st.session_state.llm_used = False

    st.session_state.ocr_lines = lines
    st.session_state.step = "review"
    st.rerun()

# ── STEP: HITL REVIEW ─────────────────────────────────────────────────────────
elif st.session_state.step == "review":
    st.markdown('<div class="step-badge">STEP 04 / HUMAN-IN-THE-LOOP REVIEW</div>', unsafe_allow_html=True)

    lines = st.session_state.ocr_lines
    flagged_count = sum(1 for l in lines if l["flagged"])
    total = len(lines)

    # Metrics row
    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.markdown(f'<div class="metric-card"><div class="metric-num">{total}</div><div class="metric-label">Lines Extracted</div></div>', unsafe_allow_html=True)
    with m2:
        st.markdown(f'<div class="metric-card"><div class="metric-num">{flagged_count}</div><div class="metric-label">LLM-Changed Lines</div></div>', unsafe_allow_html=True)
    with m3:
        heading_count = sum(1 for l in lines if l["is_heading"])
        st.markdown(f'<div class="metric-card"><div class="metric-num">{heading_count}</div><div class="metric-label">Headings Detected</div></div>', unsafe_allow_html=True)
    with m4:
        llm_label = "✓ Active" if st.session_state.llm_used else "✗ Skipped"
        st.markdown(f'<div class="metric-card"><div class="metric-num" style="font-size:1.2rem;">{llm_label}</div><div class="metric-label">Groq LLM</div></div>', unsafe_allow_html=True)

    st.markdown("---")

    if flagged_count:
        st.warning(f"⚠ **{flagged_count} line(s)** were changed by the LLM — highlighted below. Review carefully.", icon="⚠")

    st.markdown("#### ✏ Edit Transcription")
    st.caption("Correct any errors below. Heading toggle applies **bold + large font** in the Word doc.")

    edited_lines = []
    heading_vals = []

    # Column headers
    hcol1, hcol2, hcol3, hcol4 = st.columns([3, 3, 1, 1])
    hcol1.markdown("**Original (Azure)**")
    hcol2.markdown("**Agent Output — edit here**")
    hcol3.markdown("**Heading?**")
    hcol4.markdown("**Flagged**")

    for i, item in enumerate(lines):
        c1, c2, c3, c4 = st.columns([3, 3, 1, 1])
        with c1:
            st.text_input("", value=item["original_text"], disabled=True,
                          key=f"orig_{i}", label_visibility="collapsed")
        with c2:
            edited = st.text_input("", value=item["agent_text"],
                                   key=f"edit_{i}", label_visibility="collapsed")
            edited_lines.append(edited)
        with c3:
            is_h = st.checkbox("", value=item["is_heading"], key=f"head_{i}", label_visibility="collapsed")
            heading_vals.append(is_h)
        with c4:
            if item["flagged"]:
                st.markdown("🟡")
            else:
                st.markdown("·")

    st.markdown("---")
    col_btn1, col_btn2, _ = st.columns([2, 2, 4])

    with col_btn1:
        if st.button("✅  Approve & Generate Document", type="primary", use_container_width=True):
            learned = 0
            for i, item in enumerate(lines):
                item["final_text"] = edited_lines[i].strip()
                item["is_heading"] = heading_vals[i]

                # Word-level learning
                agent_words = item["agent_text"].split()
                human_words = edited_lines[i].split()
                if len(agent_words) == len(human_words) and item["agent_text"] != edited_lines[i]:
                    for aw, hw in zip(agent_words, human_words):
                        if aw != hw:
                            learn_correction(st.session_state.vocab, aw, hw)
                            audit_log("LEARN", f"Correction learned: '{aw}' → '{hw}'")
                            learned += 1

            audit_log("LEARN", f"HITL complete. {learned} correction(s) learned.")
            st.session_state.ocr_lines = lines
            st.session_state.step = "done"
            st.rerun()

    with col_btn2:
        if st.button("↩  Start Over", use_container_width=True):
            for key in ["ocr_lines", "img_width", "step", "llm_used", "consented",
                        "uploaded_file", "uploaded_name"]:
                st.session_state.pop(key, None)
            st.session_state.step = "upload"
            st.rerun()

# ── STEP: DONE / DOWNLOAD ─────────────────────────────────────────────────────
elif st.session_state.step == "done":
    st.markdown('<div class="step-badge">STEP 05 / COMPLETE</div>', unsafe_allow_html=True)
    st.success("🎉 Agent workflow complete! Your document is ready.", icon="✅")

    lines    = st.session_state.ocr_lines
    img_width = st.session_state.img_width

    with st.spinner("Generating Word document..."):
        docx_buf = generate_docx(lines, img_width)
        audit_log("ACT", "Word document generated and offered for download.")

    fname = st.session_state.get("uploaded_name", "scan")
    fname = os.path.splitext(fname)[0] + "_Agent_Processed.docx"

    st.download_button(
        label="⬇  Download Word Document",
        data=docx_buf,
        file_name=fname,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=False,
    )

    # Summary
    total   = len(lines)
    headings = sum(1 for l in lines if l["is_heading"])
    learned  = len(st.session_state.vocab)

    st.markdown("---")
    st.markdown("#### Session Summary")
    c1, c2, c3 = st.columns(3)
    c1.metric("Lines Processed", total)
    c2.metric("Headings", headings)
    c3.metric("Total Corrections Learned", learned)

    st.markdown("#### Final Transcription Preview")
    for line in lines:
        text = line["final_text"].strip()
        if not text:
            continue
        if line["is_heading"]:
            st.markdown(f"### {text}")
        else:
            st.markdown(text)

    st.markdown("---")
    if st.button("🔄  Process Another Image", type="primary"):
        for key in ["ocr_lines", "img_width", "step", "llm_used", "consented",
                    "uploaded_file", "uploaded_name"]:
            st.session_state.pop(key, None)
        st.session_state.step = "upload"
        st.rerun()
